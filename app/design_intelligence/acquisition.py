from __future__ import annotations

import hashlib
import io
from dataclasses import dataclass
from datetime import date
from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader

from .models import DesignDocumentInput, DesignProvenance

SUPPORTED_FORMATS = {
    ".md": "MARKDOWN",
    ".docx": "DOCX",
    ".pdf": "PDF",
    ".txt": "PLAIN_TEXT",
}


@dataclass(frozen=True)
class AcquisitionMetadata:
    logical_key: str
    title: str
    authors: tuple[str, ...]
    publisher: str
    source_uri: str
    license: str
    publication_date: date | None
    revision_id: int
    extraction_run_id: int
    anchor_ids: tuple[int, ...]
    evidence_link_ids: tuple[int, ...] = ()
    display_policy: str = "INTERNAL_RESEARCH_ONLY"
    topics: tuple[str, ...] = ()


class DesignDocumentAcquirer:
    """Converts authorized source bytes into an immutable BUILD-089A document input."""

    def acquire(
        self, name: str, payload: bytes, metadata: AcquisitionMetadata
    ) -> DesignDocumentInput:
        suffix = Path(name).suffix.casefold()
        if suffix not in SUPPORTED_FORMATS:
            raise ValueError("UNSUPPORTED_DESIGN_DOCUMENT_FORMAT")
        content = self._extract(suffix, payload)
        if not content.strip():
            raise ValueError("EMPTY_DESIGN_DOCUMENT")
        content = content.replace("\r\n", "\n").replace("\r", "\n")
        digest = hashlib.sha256(content.encode()).hexdigest()
        return DesignDocumentInput(
            logical_key=metadata.logical_key,
            title=metadata.title,
            content=content,
            document_type=SUPPORTED_FORMATS[suffix],
            authors=metadata.authors,
            publication_date=metadata.publication_date,
            license_metadata={
                "license": metadata.license,
                "display": metadata.display_policy,
            },
            provenance=DesignProvenance(
                source_system="design-intelligence-acquisition",
                source_id=metadata.source_uri,
                revision_id=metadata.revision_id,
                extraction_run_id=metadata.extraction_run_id,
                anchor_ids=metadata.anchor_ids,
                content_hash=digest,
                evidence_link_ids=metadata.evidence_link_ids,
            ),
            topics=metadata.topics,
            source_metadata={
                "publisher": metadata.publisher,
                "source_uri": metadata.source_uri,
                "source_filename": name,
                "acquisition_format": SUPPORTED_FORMATS[suffix],
            },
        )

    @staticmethod
    def _extract(suffix: str, payload: bytes) -> str:
        if suffix in {".md", ".txt"}:
            return payload.decode("utf-8-sig")
        if suffix == ".docx":
            document = DocxDocument(io.BytesIO(payload))
            blocks = [paragraph.text for paragraph in document.paragraphs]
            for table in document.tables:
                blocks.extend(
                    " | ".join(cell.text for cell in row.cells) for row in table.rows
                )
            return "\n".join(blocks)
        reader = PdfReader(io.BytesIO(payload))
        return "\n\f\n".join(page.extract_text() or "" for page in reader.pages)
