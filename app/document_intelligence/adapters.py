from __future__ import annotations
import csv, html, io, re, zipfile
from typing import Protocol
from .models import Anchor, Block, IntermediateDocument

class ExtractionAdapter(Protocol):
    version:str
    def extract(self,content:bytes)->IntermediateDocument: ...

def _text_document(text:str,method="VERBATIM"):
    blocks=[]; offset=0
    for seq,line in enumerate(text.splitlines(True)):
        clean=line.rstrip("\r\n"); level=None; kind="paragraph"
        m=re.match(r"^(#{1,6})\s+(.+)$",clean)
        if m: level=len(m.group(1)); clean=m.group(2); kind="heading"
        blocks.append(Block(kind,clean,seq,Anchor(char_start=offset,char_end=offset+len(line)),level))
        offset+=len(line)
    return IntermediateDocument(({"logical_unit":1,"text":text},),tuple(blocks),extraction_method=method)

class PlainTextAdapter:
    version="plain-1"
    def extract(self,content): return _text_document(content.decode("utf-8",errors="replace"))
class MarkdownAdapter(PlainTextAdapter): version="markdown-1"
class HtmlAdapter:
    version="html-1"
    def extract(self,content):
        text=html.unescape(re.sub(r"<[^>]+>","\n",content.decode("utf-8",errors="replace")))
        return _text_document(text,"NORMALIZED_HTML")
class CsvAdapter:
    version="csv-1"
    def extract(self,content):
        text=content.decode("utf-8-sig",errors="replace"); rows=list(csv.reader(io.StringIO(text)))
        return IntermediateDocument(({"logical_unit":1,"text":text},),(Block("table","",0,Anchor(char_start=0,char_end=len(text))),),tables=({"rows":rows,"complete":True},))

class PdfAdapter:
    version="pypdf-1"
    def extract(self,content):
        from pypdf import PdfReader
        pages=[]; blocks=[]; warnings=[]; offset=sequence=0
        try: reader=PdfReader(io.BytesIO(content),strict=False)
        except Exception as exc: return IntermediateDocument((),(),warnings=(f"PDF_ADAPTER_FAILURE:{exc.__class__.__name__}",))
        for number,page in enumerate(reader.pages,1):
            try: text=page.extract_text() or ""
            except Exception as exc: text=""; warnings.append(f"PAGE_{number}_EXTRACTION_FAILED:{exc.__class__.__name__}")
            method="EMBEDDED_TEXT" if text.strip() else "OCR_REQUIRED"
            if method=="OCR_REQUIRED": warnings.append(f"PAGE_{number}_OCR_REQUIRED")
            pages.append({"page":number,"text":text,"method":method})
            for line in text.splitlines(True):
                clean=line.strip(); kind="caption" if re.match(r"^(fig(?:ure)?|table)\s*\d+",clean,re.I) else "paragraph"
                if clean and (clean.isupper() or (len(clean)<100 and clean.endswith(":"))): kind="heading"
                blocks.append(Block(kind,clean,sequence,Anchor(number,number,offset,offset+len(line),{"page":number}),1 if kind=="heading" else None)); sequence+=1; offset+=len(line)
        return IntermediateDocument(tuple(pages),tuple(blocks),warnings=tuple(warnings),extraction_method="MIXED" if warnings else "VERBATIM")

class DocxAdapter:
    version="python-docx-1"
    def extract(self,content):
        from docx import Document
        document=Document(io.BytesIO(content)); blocks=[]; tables=[]; media=[]; sequence=offset=0
        for paragraph in document.paragraphs:
            text=paragraph.text; style=paragraph.style.name if paragraph.style else ""; match=re.match(r"Heading\s+(\d+)",style,re.I); level=int(match.group(1)) if match else None
            kind="heading" if level else ("caption" if style.casefold()=="caption" else ("list" if style.startswith("List") else "paragraph"))
            blocks.append(Block(kind,text,sequence,Anchor(char_start=offset,char_end=offset+len(text)),level,{"style":style})); sequence+=1; offset+=len(text)+1
        for number,table in enumerate(document.tables,1):
            tables.append({"table":number,"rows":[[c.text for c in row.cells] for row in table.rows],"complete":True}); blocks.append(Block("table",f"Table {number}",sequence,Anchor(),metadata={"table":number})); sequence+=1
        with zipfile.ZipFile(io.BytesIO(content)) as archive:
            media.extend({"package_path":name,"kind":"embedded_media"} for name in archive.namelist() if name.startswith("word/media/"))
        return IntermediateDocument(({"logical_unit":1,"text":"\n".join(b.text for b in blocks)},),tuple(blocks),tuple(tables),tuple(media))

class GoogleDocsExportAdapter(DocxAdapter):
    version="google-docs-docx-export-1"
    def extract(self,content):
        value=super().extract(content); return IntermediateDocument(value.units,value.blocks,value.tables,value.media,value.warnings,"GOOGLE_DOCS_EXPORTED_VERBATIM")

ADAPTERS={"application/pdf":PdfAdapter(),"application/vnd.openxmlformats-officedocument.wordprocessingml.document":DocxAdapter(),"application/vnd.google-apps.document":GoogleDocsExportAdapter(),"text/plain":PlainTextAdapter(),"text/markdown":MarkdownAdapter(),"text/x-markdown":MarkdownAdapter(),"text/html":HtmlAdapter(),"text/csv":CsvAdapter()}
