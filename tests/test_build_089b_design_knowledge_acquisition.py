import hashlib
import io
import os
from datetime import date
from pathlib import Path

import pytest
from docx import Document as DocxDocument
from fastapi import FastAPI
from fastapi.testclient import TestClient

from app.design_intelligence import (
    AcquisitionMetadata,
    DesignDocument,
    DesignDocumentAcquirer,
    DesignProvenance,
    DesignReasoningService,
    EducationalClassification,
    MemoryDesignKnowledgeRepository,
    SemanticDesignDomain,
    SemanticUnitType,
)
from app.design_intelligence.models import utcnow
from app.security import verify_owner_or_api_key


def metadata() -> AcquisitionMetadata:
    return AcquisitionMetadata(
        logical_key="dashboard-learning",
        title="Dashboard learning guidance",
        authors=("Design Researcher",),
        publisher="Orchid University",
        source_uri="https://example.test/design/dashboard",
        license="CC-BY-4.0",
        publication_date=date(2025, 3, 1),
        revision_id=501,
        extraction_run_id=601,
        anchor_ids=(701,),
        evidence_link_ids=(801,),
        display_policy="FULL_TEXT_ALLOWED",
        topics=("dashboard", "education"),
    )


def design_document(content: str, document_id: int = 1) -> DesignDocument:
    return DesignDocument(
        document_id=document_id,
        logical_key=f"design-{document_id}",
        version=1,
        title="Design guidance",
        content=content,
        document_type="MARKDOWN",
        authors=("Researcher",),
        publication_date=date(2025, 1, 1),
        license_metadata={"license": "CC-BY-4.0", "display": "FULL_TEXT_ALLOWED"},
        provenance=DesignProvenance(
            source_system="test",
            source_id=f"source:{document_id}",
            revision_id=501,
            extraction_run_id=601,
            anchor_ids=(701,),
            content_hash=hashlib.sha256(content.encode()).hexdigest(),
            evidence_link_ids=(801,),
        ),
        domains=(),
        knowledge_types=(),
        topics=("design",),
        classification_confidence=0.8,
        classification_version="089a-design-rules-1",
        source_metadata={"publisher": "Test"},
        created_at=utcnow(),
    )


def test_acquisition_supports_markdown_text_and_docx_with_immutable_metadata():
    acquirer = DesignDocumentAcquirer()
    markdown = acquirer.acquire(
        "guide.md", b"# Guide\nDashboard guideline.", metadata()
    )
    assert markdown.document_type == "MARKDOWN"
    assert markdown.provenance.source_id == "https://example.test/design/dashboard"
    assert markdown.source_metadata["publisher"] == "Orchid University"
    assert markdown.license_metadata["license"] == "CC-BY-4.0"
    plain = acquirer.acquire("guide.txt", b"Accessibility standard.", metadata())
    assert plain.document_type == "PLAIN_TEXT"
    document = DocxDocument()
    document.add_heading("Motion", 1)
    document.add_paragraph("Animation should honor reduced motion.")
    payload = io.BytesIO()
    document.save(payload)
    docx = acquirer.acquire("guide.docx", payload.getvalue(), metadata())
    assert docx.document_type == "DOCX" and "reduced motion" in docx.content
    with pytest.raises(ValueError, match="UNSUPPORTED"):
        acquirer.acquire("guide.html", b"<p>no</p>", metadata())


def test_decomposition_preserves_structure_parentage_and_exact_source_locations():
    content = """# Dashboard principles
Dashboard guidance communicates status.
- Use progressive disclosure.
1. Identify user goals.
| Metric | Meaning |
> Prefer direct labels.
Warning: animation must honor reduced motion.
```
Button(label="Save")
```
"""
    units = DesignReasoningService().index_document(design_document(content))
    kinds = {unit.unit_type for unit in units}
    assert {
        SemanticUnitType.HEADING,
        SemanticUnitType.PARAGRAPH,
        SemanticUnitType.BULLET_LIST,
        SemanticUnitType.NUMBERED_PROCEDURE,
        SemanticUnitType.TABLE,
        SemanticUnitType.QUOTED_GUIDANCE,
        SemanticUnitType.WARNING,
        SemanticUnitType.CODE_EXAMPLE,
    } <= kinds
    assert all(unit.source_location.start <= unit.source_location.end for unit in units)
    assert all(unit.source_location.locator["anchor_ids"] == [701] for unit in units)
    assert all(unit.parent_unit_id == units[0].unit_id for unit in units[1:])


def test_multilabel_classification_embeddings_graph_and_idempotency():
    repository = MemoryDesignKnowledgeRepository()
    service = DesignReasoningService(repository)
    content = """# Learning dashboard
Mayer multimedia learning and cognitive load theory support dashboard accessibility guidelines.
Accessibility requires keyboard interaction and WCAG conformance.
Dashboard visualization best practice improves scientific communication.
"""
    first = service.index_document(design_document(content))
    second = service.index_document(design_document(content))
    assert [unit.unit_id for unit in first] == [unit.unit_id for unit in second]
    assert len(repository.units) == len(first)
    classified = next(unit for unit in first if "Mayer" in unit.text)
    assert SemanticDesignDomain.DASHBOARD_DESIGN in classified.domains
    assert SemanticDesignDomain.ACCESSIBILITY in classified.domains
    assert (
        EducationalClassification.MAYER_MULTIMEDIA_LEARNING
        in classified.educational_classifications
    )
    assert (
        EducationalClassification.COGNITIVE_LOAD_THEORY
        in classified.educational_classifications
    )
    assert len(classified.embedding) == 32
    assert classified.embedding_metadata["local_execution"] is True
    assert repository.relationships and all(
        rel.provenance for rel in repository.relationships
    )
    assert len(repository.audit_events) == len(repository.units) + len(
        repository.relationships
    )


def test_hybrid_reasoning_retrieval_returns_explanations_citations_and_related_concepts():
    service = DesignReasoningService()
    service.index_document(
        design_document(
            "# Dashboard\nDashboard accessibility guideline should use keyboard interaction.\n"
            "This design principle supports scientific visualization.",
            10,
        )
    )
    result = service.search(
        "Find dashboard accessibility guidance",
        domains=(SemanticDesignDomain.DASHBOARD_DESIGN,),
    )
    assert result["total"] >= 1
    item = result["results"][0]
    assert 0 < item["confidence"] <= 1
    assert item["supporting_citations"][0]["locator"]["anchor_ids"] == [701]
    assert item["classification"]["domains"]
    assert item["explanation"]["formula"]
    assert item["related_concepts"]


def test_authenticated_reasoning_route_remains_read_only(monkeypatch):
    from app.design_intelligence import routes

    service = DesignReasoningService()
    service.index_document(
        design_document("Dashboard guideline should use clear status labels.")
    )
    monkeypatch.setattr(routes, "REASONING_SERVICE", service)
    app = FastAPI()
    app.include_router(routes.router)
    app.dependency_overrides[verify_owner_or_api_key] = lambda: {"actor": "test"}
    client = TestClient(app)
    assert (
        client.post(
            "/api/design-intelligence/reasoning-search",
            json={"query": "dashboard guidance"},
        ).status_code
        == 200
    )
    assert client.get("/api/design-intelligence/reasoning-search").status_code == 405
    configuration = client.get("/api/design-intelligence/configuration").json()
    assert (
        configuration["read_only"] is True
        and len(configuration["semantic_domains"]) == 18
    )


def test_089b_migration_is_additive_indexed_append_only_and_provenance_preserving():
    sql = Path("migrations/089b_design_knowledge_acquisition.sql").read_text()
    assert "DROP TABLE" not in sql and "TRUNCATE" not in sql
    assert "REFERENCES oc_design_intelligence.documents" in sql
    assert "source_location JSONB NOT NULL" in sql
    assert "semantic_relationships" in sql and "protect_089b_" in sql
    assert "design_semantic_fts_idx" in sql


@pytest.mark.skipif(
    not os.getenv("TEST_DATABASE_URL"), reason="TEST_DATABASE_URL not configured"
)
def test_postgresql_authoritative_units_relationships_and_append_only_guards():
    import psycopg
    from app.design_intelligence.postgres_knowledge_repository import (
        PostgresDesignKnowledgeRepository,
    )

    dsn = os.environ["TEST_DATABASE_URL"]
    with psycopg.connect(dsn) as con:
        document_id = con.execute(
            "SELECT document_id FROM oc_design_intelligence.documents ORDER BY document_id LIMIT 1"
        ).fetchone()[0]
    repository = PostgresDesignKnowledgeRepository(dsn)
    service = DesignReasoningService(repository)
    units = service.index_document(
        design_document(
            "Dashboard guideline should support accessibility.", document_id
        )
    )
    assert repository.units[-1].unit_id == units[-1].unit_id
    with psycopg.connect(dsn) as con, pytest.raises(psycopg.Error, match="APPEND_ONLY"):
        con.execute(
            "UPDATE oc_design_intelligence.semantic_units SET authorized_text='changed' WHERE unit_id=%s",
            (units[0].unit_id,),
        )
