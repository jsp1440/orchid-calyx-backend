from io import BytesIO
from pathlib import Path

from docx import Document
from pypdf import PdfWriter

from app.document_intelligence.adapters import DocxAdapter, GoogleDocsExportAdapter, PdfAdapter
from app.document_intelligence.lifecycle import ExtractionService, TRANSITIONS
from app.document_intelligence.memory_repository import MemoryRepository


def source(revision=1, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", content=None):
    if content is None:
        d=Document(); d.add_heading("Methods",1); d.add_paragraph("First paragraph"); d.add_table(1,1).cell(0,0).text="datum"; b=BytesIO(); d.save(b); content=b.getvalue()
    return {"revision_id":revision,"registry_id":revision,"sha256":"a"*64,"filename":"study.docx","mime_type":mime,"content":content,"provenance":{"drive_revision":"immutable"}}


def test_migration_is_additive_and_evidence_linked():
    sql=Path("migrations/084_document_intelligence.sql").read_text()
    assert "CREATE SCHEMA IF NOT EXISTS oc_document_intelligence" in sql
    assert "REFERENCES oc_import.document_revisions" in sql
    assert "configuration_hash" in sql and "DERIVING_OBJECTS" in sql
    assert all(x not in sql.upper() for x in ("DROP ","TRUNCATE ","ALTER SCHEMA"))


def test_lifecycle_idempotency_history_and_versioning():
    repo=MemoryRepository([source()]); svc=ExtractionService(repo)
    first=svc.start(1,"extractor-1","rules-1",{"layout":True})
    same=svc.start(1,"extractor-1","rules-1",{"layout":True})
    changed=svc.start(1,"extractor-2","rules-1",{"layout":True})
    assert first["state"] == same["state"] == changed["state"] == "COMPLETED"
    assert first["extraction_run_id"] == same["extraction_run_id"]
    assert first["extraction_run_id"] != changed["extraction_run_id"]
    assert len(repo.history(first["record_id"])) == 2
    assert "DERIVING_OBJECTS" in TRANSITIONS["STRUCTURED"]


def test_safe_cancellation():
    repo=MemoryRepository([source()]); record=repo.ensure_record(source())
    run=repo.ensure_run(record["record_id"],"v","r","0"*64)
    repo.request_cancel(run["extraction_run_id"])
    assert ExtractionService(repo).resume(run["extraction_run_id"])["state"] == "CANCELLED"


def test_pdf_boundaries_and_image_only_warnings():
    out=BytesIO(); w=PdfWriter(); w.add_blank_page(72,72); w.add_blank_page(72,72); w.write(out)
    doc=PdfAdapter().extract(out.getvalue())
    assert [u["page"] for u in doc.units] == [1,2]
    assert len(doc.warnings)==2 and all("OCR_REQUIRED" in warning for warning in doc.warnings)


def test_docx_and_google_docs_structure_and_order():
    s=source(); doc=DocxAdapter().extract(s["content"])
    assert doc.blocks[0].kind == "heading" and doc.blocks[0].heading_level == 1
    assert any(b.kind == "paragraph" for b in doc.blocks) and doc.tables
    google=GoogleDocsExportAdapter().extract(s["content"])
    assert google.extraction_method == "GOOGLE_DOCS_EXPORTED_VERBATIM"


def test_derived_chunks_keep_complete_parent_pointer_and_anchors():
    repo=MemoryRepository([source()]); run=ExtractionService(repo).start(1,"v","r")
    assert any(s["run_id"] == run["extraction_run_id"] for s in repo.structures)
    assert all(c["complete_object_pointer"]["structural_id"] for c in repo.chunks)
    assert all(a["revision_id"] == 1 for a in repo.anchors)


def test_no_drive_writes_or_graph_publication():
    code="\n".join(p.read_text() for p in Path("app/document_intelligence").glob("*.py"))
    assert "drive.files.update" not in code and "drive.files.delete" not in code
    assert "knowledge_graph" not in code and "embedding" not in code
